"""
Microsoft Presidio 기반 PII / 민감정보 감지 로컬 웹 서버.

실행:
    python app.py
브라우저에서 http://127.0.0.1:5000 접속 후 파일을 드래그-앤-드롭하면
감지된 엔티티 목록과 위치, 신뢰도, 매치 텍스트를 확인할 수 있습니다.

custom_patterns.yaml 을 수정하여 사용자 정의 패턴/룰을 자유롭게 추가할 수 있습니다.
"""

from __future__ import annotations

import io
import logging
import os
from pathlib import Path
from typing import List

import yaml
from flask import Flask, jsonify, render_template, request
from presidio_analyzer import (
    AnalyzerEngine,
    Pattern,
    PatternRecognizer,
    RecognizerRegistry,
)
from presidio_analyzer.nlp_engine import NlpEngineProvider

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger("presidio-app")

BASE_DIR = Path(__file__).resolve().parent
PATTERN_FILE = BASE_DIR / "custom_patterns.yaml"

# ---------------------------------------------------------------------------
# Presidio Analyzer 초기화 (en + ko 지원)
# ---------------------------------------------------------------------------
# 사용 가능한 spaCy 모델을 자동 감지해 NLP 엔진을 구성합니다.
# - 영문: en_core_web_sm (또는 _md/_lg)
# - 한국어: ko_core_news_sm (또는 _md/_lg)
# 한국어 모델 설치:  python -m spacy download ko_core_news_sm
# 영문 모델 설치 :  python -m spacy download en_core_web_sm

_LANG_MODEL_CANDIDATES = [
    ("en", ["en_core_web_lg", "en_core_web_md", "en_core_web_sm"]),
    ("ko", ["ko_core_news_lg", "ko_core_news_md", "ko_core_news_sm"]),
]


def _detect_models() -> dict:
    """설치된 spaCy 모델만 lang_code 별로 1개씩 골라 반환."""
    import spacy
    chosen: dict = {}
    for lang, candidates in _LANG_MODEL_CANDIDATES:
        for model in candidates:
            try:
                spacy.load(model)
                chosen[lang] = model
                log.info("spaCy: %s = %s", lang, model)
                break
            except OSError:
                continue
        if lang not in chosen:
            log.warning("spaCy: %s 모델 미설치 → %s 비활성화", lang, lang)
    if not chosen:
        raise RuntimeError(
            "사용 가능한 spaCy 모델이 없습니다. "
            "python -m spacy download en_core_web_sm 또는 ko_core_news_sm 으로 설치하세요."
        )
    return chosen


_MODELS = _detect_models()
SUPPORTED_LANGUAGES = list(_MODELS.keys())  # 예: ["en", "ko"]
NLP_CONFIG = {
    "nlp_engine_name": "spacy",
    "models": [{"lang_code": k, "model_name": v} for k, v in _MODELS.items()],
}


def _normalize_languages(spec) -> List[str]:
    """YAML 의 supported_language → 실제 적용 언어 리스트.
    'any' / None → 설치된 모든 언어. 문자열/리스트는 그대로 (단 미지원 언어는 제거)."""
    if spec is None or spec == "any":
        return list(SUPPORTED_LANGUAGES)
    if isinstance(spec, str):
        spec = [spec]
    return [s for s in spec if s in SUPPORTED_LANGUAGES] or [SUPPORTED_LANGUAGES[0]]


def load_custom_recognizers(path: Path) -> List[PatternRecognizer]:
    """custom_patterns.yaml 을 읽어 PatternRecognizer 목록을 만든다.
    supported_language 가 'any' 또는 리스트면 언어별로 복제해 등록한다."""
    if not path.exists():
        log.warning("custom_patterns.yaml not found: %s", path)
        return []

    with path.open("r", encoding="utf-8") as f:
        cfg = yaml.safe_load(f) or {}

    recognizers: List[PatternRecognizer] = []

    for item in cfg.get("pattern_recognizers", []) or []:
        patterns = [
            Pattern(name=p["name"], regex=p["regex"], score=float(p["score"]))
            for p in item.get("patterns", [])
        ]
        for lang in _normalize_languages(item.get("supported_language")):
            rec = PatternRecognizer(
                supported_entity=item["supported_entity"],
                name=f"{item['name']}_{lang}" if len(SUPPORTED_LANGUAGES) > 1 else item["name"],
                supported_language=lang,
                patterns=patterns,
                context=item.get("context"),
            )
            recognizers.append(rec)
        log.info("Loaded pattern recognizer: %s (%d patterns, langs=%s)",
                 item["name"], len(patterns), _normalize_languages(item.get("supported_language")))

    for item in cfg.get("deny_list_recognizers", []) or []:
        for lang in _normalize_languages(item.get("supported_language")):
            rec = PatternRecognizer(
                supported_entity=item["supported_entity"],
                name=f"{item['name']}_{lang}" if len(SUPPORTED_LANGUAGES) > 1 else item["name"],
                supported_language=lang,
                deny_list=item.get("deny_list", []),
                deny_list_score=float(item.get("score", 1.0)),
            )
            recognizers.append(rec)
        log.info(
            "Loaded deny-list recognizer: %s (%d terms, langs=%s)",
            item["name"], len(item.get("deny_list", [])),
            _normalize_languages(item.get("supported_language")),
        )

    return recognizers


def build_analyzer() -> AnalyzerEngine:
    """기본 Presidio 인식기 + custom_patterns.yaml 의 사용자 정의 인식기를 합쳐 엔진 생성."""
    provider = NlpEngineProvider(nlp_configuration=NLP_CONFIG)
    nlp_engine = provider.create_engine()

    registry = RecognizerRegistry()
    # Presidio 빌트인 인식기는 영문 한정 — 한국어 NLP 엔진에서는 일부만 호환
    registry.load_predefined_recognizers(nlp_engine=nlp_engine, languages=SUPPORTED_LANGUAGES)

    for rec in load_custom_recognizers(PATTERN_FILE):
        registry.add_recognizer(rec)

    return AnalyzerEngine(
        registry=registry,
        nlp_engine=nlp_engine,
        supported_languages=SUPPORTED_LANGUAGES,
    )


analyzer: AnalyzerEngine = build_analyzer()

# ---------------------------------------------------------------------------
# 파일 → 텍스트 추출
# ---------------------------------------------------------------------------

def _smart_decode(raw: bytes) -> str:
    """BOM 우선 → utf-8 strict → cp949 / euc-kr → utf-8 replace."""
    if raw.startswith(b"\xef\xbb\xbf"):
        return raw[3:].decode("utf-8", errors="replace")
    if raw.startswith(b"\xff\xfe"):
        return raw[2:].decode("utf-16-le", errors="replace")
    if raw.startswith(b"\xfe\xff"):
        return raw[2:].decode("utf-16-be", errors="replace")
    for enc in ("utf-8", "cp949", "euc-kr"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def extract_text(filename: str, raw: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext in {".txt", ".log", ".csv", ".json", ".md", ".yaml", ".yml", ".xml", ".html"}:
        return _smart_decode(raw)

    if ext == ".pdf":
        from pypdf import PdfReader  # lazy import
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join((p.extract_text() or "") for p in reader.pages)

    if ext == ".docx":
        from docx import Document  # lazy import
        doc = Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    return _smart_decode(raw)


# ---------------------------------------------------------------------------
# Flask
# ---------------------------------------------------------------------------

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 32 * 1024 * 1024  # 32MB


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/docs")
def docs():
    return render_template("docs.html")


@app.route("/api/recognizers", methods=["GET"])
def list_recognizers():
    """현재 활성화된 인식기 목록과 지원 엔티티를 반환."""
    recs = analyzer.registry.recognizers
    out = []
    for r in recs:
        out.append(
            {
                "name": r.name,
                "supported_entities": list(r.supported_entities),
                "supported_language": r.supported_language,
            }
        )
    return jsonify(
        {
            "recognizer_count": len(out),
            "supported_entities": sorted(
                {e for r in recs for e in r.supported_entities}
            ),
            "recognizers": out,
        }
    )


@app.route("/api/reload", methods=["POST"])
def reload_patterns():
    """custom_patterns.yaml 을 다시 읽어 분석기를 재구성한다."""
    global analyzer
    try:
        analyzer = build_analyzer()
        return jsonify({"ok": True, "message": "Patterns reloaded."})
    except Exception as e:  # noqa: BLE001
        log.exception("reload failed")
        return jsonify({"ok": False, "message": str(e)}), 500


@app.route("/api/analyze", methods=["POST"])
def analyze():
    if "file" not in request.files:
        return jsonify({"ok": False, "message": "file field is required"}), 400

    f = request.files["file"]
    raw = f.read()
    score_threshold = float(request.form.get("score_threshold", 0.3))
    language = (request.form.get("language") or "auto").strip().lower()

    try:
        text = extract_text(f.filename or "uploaded", raw)
    except Exception as e:  # noqa: BLE001
        log.exception("extract_text failed")
        return jsonify({"ok": False, "message": f"failed to read file: {e}"}), 400

    # auto = 설치된 모든 언어로 분석 후 (start, end, entity_type) 으로 dedup
    if language == "auto" or language not in SUPPORTED_LANGUAGES:
        target_langs = list(SUPPORTED_LANGUAGES)
    else:
        target_langs = [language]

    raw_results = []
    for lang in target_langs:
        try:
            raw_results.extend(analyzer.analyze(
                text=text, language=lang, score_threshold=score_threshold,
            ))
        except Exception as e:  # noqa: BLE001
            log.warning("analyze(%s) failed: %s", lang, e)

    # dedup: 같은 위치/엔티티 중 최고 score 만 유지
    dedup: dict = {}
    for r in raw_results:
        k = (r.start, r.end, r.entity_type)
        if k not in dedup or dedup[k].score < r.score:
            dedup[k] = r
    results = list(dedup.values())

    findings = []
    for r in results:
        snippet = text[r.start : r.end]
        findings.append(
            {
                "entity_type": r.entity_type,
                "start": r.start,
                "end": r.end,
                "score": round(float(r.score), 3),
                "text": snippet,
                "recognizer": (
                    r.recognition_metadata.get("recognizer_name")
                    if r.recognition_metadata
                    else None
                ),
            }
        )

    findings.sort(key=lambda x: (x["start"], -x["score"]))

    summary: dict[str, int] = {}
    for f_ in findings:
        summary[f_["entity_type"]] = summary.get(f_["entity_type"], 0) + 1

    return jsonify(
        {
            "ok": True,
            "filename": f.filename,
            "char_count": len(text),
            "score_threshold": score_threshold,
            "language": language if language in ("auto",) + tuple(SUPPORTED_LANGUAGES) else "auto",
            "languages_used": target_langs,
            "summary": summary,
            "findings": findings,
            "text": text,
        }
    )


if __name__ == "__main__":
    host = os.environ.get("HOST", "127.0.0.1")
    port = int(os.environ.get("PORT", "5000"))
    log.info("Starting server on http://%s:%s", host, port)
    app.run(host=host, port=port, debug=False)
