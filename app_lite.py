"""
Presidio/spaCy 의존성 없는 lite 버전.

- Microsoft Presidio 의 PatternRecognizer 동작(정규식 + context boost + deny-list)을
  순수 Python 으로 재구현.
- 같은 custom_patterns.yaml 포맷 사용, 같은 UI(templates/index.html) 사용.
- Python 3.14 에서도 추가 컴파일 없이 그대로 동작.

실행:  python app_lite.py
"""

from __future__ import annotations

import io
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

import yaml
from flask import Flask, jsonify, render_template, request

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger("presidio-lite")

BASE_DIR = Path(__file__).resolve().parent
PATTERN_FILE = BASE_DIR / "custom_patterns.yaml"

# ---------------------------------------------------------------------------
# Pattern recognizer (Presidio 동작 모사)
# ---------------------------------------------------------------------------

CONTEXT_WINDOW = 60   # context 단어를 매치 주변 ±N 글자에서 탐색
CONTEXT_BOOST = 0.35  # context 단어가 있을 때 score 가산치 (cap 1.0)


@dataclass
class CompiledPattern:
    name: str
    regex: re.Pattern
    score: float


@dataclass
class Recognizer:
    name: str
    entity: str
    patterns: List[CompiledPattern] = field(default_factory=list)
    deny_list: List[str] = field(default_factory=list)
    deny_list_score: float = 1.0
    context: List[str] = field(default_factory=list)
    languages: List[str] = field(default_factory=lambda: ["any"])

    def matches_language(self, language: str) -> bool:
        if language == "auto" or "any" in self.languages:
            return True
        return language in self.languages

    def analyze(self, text: str, lower_text: str) -> List[dict]:
        out: List[dict] = []

        for pat in self.patterns:
            for m in pat.regex.finditer(text):
                score = pat.score
                if self.context:
                    s = max(0, m.start() - CONTEXT_WINDOW)
                    e = min(len(text), m.end() + CONTEXT_WINDOW)
                    window = lower_text[s:e]
                    if any(c.lower() in window for c in self.context):
                        score = min(1.0, score + CONTEXT_BOOST)
                out.append(
                    {
                        "entity_type": self.entity,
                        "start": m.start(),
                        "end": m.end(),
                        "score": round(score, 3),
                        "text": text[m.start() : m.end()],
                        "recognizer": self.name,
                    }
                )

        for term in self.deny_list:
            if not term:
                continue
            for m in re.finditer(re.escape(term), text, flags=re.IGNORECASE):
                out.append(
                    {
                        "entity_type": self.entity,
                        "start": m.start(),
                        "end": m.end(),
                        "score": round(self.deny_list_score, 3),
                        "text": text[m.start() : m.end()],
                        "recognizer": self.name,
                    }
                )
        return out


# ---------------------------------------------------------------------------
# 빌트인 인식기 (Presidio 의 주요 정규식 룰을 옮긴 것)
# ---------------------------------------------------------------------------

def _build_builtin() -> List[Recognizer]:
    return [
        Recognizer(
            name="EmailRecognizer",
            entity="EMAIL_ADDRESS",
            patterns=[CompiledPattern(
                "email",
                re.compile(r"\b[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}\b"),
                0.9,
            )],
            context=["email", "이메일", "메일"],
        ),
        Recognizer(
            name="UrlRecognizer",
            entity="URL",
            patterns=[CompiledPattern(
                "url",
                re.compile(r"\bhttps?://[^\s<>\"']+", re.IGNORECASE),
                0.6,
            )],
        ),
        Recognizer(
            name="IpRecognizer",
            entity="IP_ADDRESS",
            patterns=[
                CompiledPattern(
                    "ipv4",
                    re.compile(r"\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d?\d)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d?\d)\b"),
                    0.6,
                ),
                CompiledPattern(
                    "ipv6",
                    re.compile(r"\b(?:[A-Fa-f0-9]{1,4}:){7}[A-Fa-f0-9]{1,4}\b"),
                    0.7,
                ),
            ],
            context=["ip", "ipv4", "ipv6", "address", "주소"],
        ),
        Recognizer(
            name="CreditCardRecognizer",
            entity="CREDIT_CARD",
            patterns=[CompiledPattern(
                "cc",
                re.compile(r"\b(?:\d[ -]*?){13,19}\b"),
                0.4,  # Luhn 검증 통과 시 후처리에서 1.0 으로 상승
            )],
            context=["card", "credit", "신용카드", "카드번호", "카드"],
        ),
        Recognizer(
            name="UsSsnRecognizer",
            entity="US_SSN",
            patterns=[CompiledPattern(
                "ssn",
                re.compile(r"\b(?!000|666|9\d{2})\d{3}[- ]?(?!00)\d{2}[- ]?(?!0000)\d{4}\b"),
                0.55,
            )],
            context=["ssn", "social", "security"],
        ),
        Recognizer(
            name="IbanRecognizer",
            entity="IBAN_CODE",
            patterns=[CompiledPattern(
                "iban",
                re.compile(r"\b[A-Z]{2}\d{2}[A-Z0-9]{11,30}\b"),
                0.5,
            )],
            context=["iban", "bank", "은행", "계좌"],
        ),
        Recognizer(
            name="UsPhoneRecognizer",
            entity="PHONE_NUMBER",
            patterns=[CompiledPattern(
                "us_phone",
                re.compile(r"\b(?:\+?1[ \-.])?\(?\d{3}\)?[ \-.]\d{3}[ \-.]\d{4}\b"),
                0.5,
            )],
            context=["phone", "tel", "전화", "연락처"],
        ),
    ]


def _normalize_languages(spec) -> List[str]:
    """Accept str / list / 'any' / None → list of languages."""
    if spec is None:
        return ["any"]
    if isinstance(spec, str):
        return [spec]
    if isinstance(spec, list):
        return [str(x) for x in spec] or ["any"]
    return ["any"]


# ---------------------------------------------------------------------------
# YAML 로드
# ---------------------------------------------------------------------------

def load_custom(path: Path) -> List[Recognizer]:
    if not path.exists():
        log.warning("custom_patterns.yaml not found: %s", path)
        return []

    with path.open("r", encoding="utf-8") as f:
        cfg = yaml.safe_load(f) or {}

    recs: List[Recognizer] = []

    for item in cfg.get("pattern_recognizers", []) or []:
        try:
            patterns = [
                CompiledPattern(
                    name=p["name"],
                    regex=re.compile(p["regex"]),
                    score=float(p["score"]),
                )
                for p in item.get("patterns", [])
            ]
        except re.error as e:
            log.error("Invalid regex in %s: %s", item.get("name"), e)
            continue

        recs.append(
            Recognizer(
                name=item["name"],
                entity=item["supported_entity"],
                patterns=patterns,
                context=item.get("context") or [],
                languages=_normalize_languages(item.get("supported_language")),
            )
        )
        log.info(
            "Loaded pattern recognizer: %s (%d patterns, langs=%s)",
            item["name"], len(patterns), recs[-1].languages,
        )

    for item in cfg.get("deny_list_recognizers", []) or []:
        recs.append(
            Recognizer(
                name=item["name"],
                entity=item["supported_entity"],
                deny_list=item.get("deny_list") or [],
                deny_list_score=float(item.get("score", 1.0)),
                languages=_normalize_languages(item.get("supported_language")),
            )
        )
        log.info(
            "Loaded deny-list recognizer: %s (%d terms, langs=%s)",
            item["name"], len(item.get("deny_list") or []), recs[-1].languages,
        )

    return recs


def build_recognizers() -> List[Recognizer]:
    return _build_builtin() + load_custom(PATTERN_FILE)


recognizers: List[Recognizer] = build_recognizers()

# ---------------------------------------------------------------------------
# Luhn 검증 (CREDIT_CARD score 보정)
# ---------------------------------------------------------------------------

def _luhn_ok(number: str) -> bool:
    digits = [int(c) for c in number if c.isdigit()]
    if not 13 <= len(digits) <= 19:
        return False
    s = 0
    for i, d in enumerate(reversed(digits)):
        if i % 2 == 1:
            d *= 2
            if d > 9:
                d -= 9
        s += d
    return s % 10 == 0


# ---------------------------------------------------------------------------
# 분석
# ---------------------------------------------------------------------------

def analyze_text(text: str, score_threshold: float, language: str = "auto") -> List[dict]:
    lower = text.lower()
    findings: List[dict] = []
    for r in recognizers:
        if not r.matches_language(language):
            continue
        findings.extend(r.analyze(text, lower))

    # CREDIT_CARD: Luhn 통과한 것만 신뢰도 상승, 실패한 건 점수 하락
    for f in findings:
        if f["entity_type"] == "CREDIT_CARD":
            f["score"] = round(0.95 if _luhn_ok(f["text"]) else max(0.0, f["score"] - 0.3), 3)

    findings = [f for f in findings if f["score"] >= score_threshold]

    # 동일 위치 + 동일 엔티티 중복 제거 (가장 높은 score 만 유지)
    dedup: dict[tuple, dict] = {}
    for f in findings:
        k = (f["start"], f["end"], f["entity_type"])
        if k not in dedup or dedup[k]["score"] < f["score"]:
            dedup[k] = f
    findings = sorted(dedup.values(), key=lambda x: (x["start"], -x["score"]))
    return findings


# ---------------------------------------------------------------------------
# 파일 → 텍스트
# ---------------------------------------------------------------------------

def _smart_decode(raw: bytes) -> str:
    """BOM 우선 → utf-8 strict → cp949 / euc-kr → utf-8 replace."""
    if raw.startswith(b"\xef\xbb\xbf"):
        return raw[3:].decode("utf-8", errors="replace")
    if raw.startswith(b"\xff\xfe"):
        return raw[2:].decode("utf-16-le", errors="replace")
    if raw.startswith(b"\xfe\xff"):
        return raw[2:].decode("utf-16-be", errors="replace")
    for enc in ("utf-8", "cp949", "euc-kr"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def extract_text(filename: str, raw: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext in {".txt", ".log", ".csv", ".json", ".md", ".yaml", ".yml", ".xml", ".html"}:
        return _smart_decode(raw)

    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join((p.extract_text() or "") for p in reader.pages)

    if ext == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    if ext == ".hwp":
        from hwp_extract import extract_hwp_text
        return extract_hwp_text(raw)

    return _smart_decode(raw)


# ---------------------------------------------------------------------------
# Flask
# ---------------------------------------------------------------------------

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 32 * 1024 * 1024


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/docs")
def docs():
    return render_template("docs.html")


@app.route("/api/recognizers", methods=["GET"])
def list_recognizers():
    out = [
        {
            "name": r.name,
            "supported_entities": [r.entity],
            "languages": r.languages,
            "supported_language": "any",
        }
        for r in recognizers
    ]
    return jsonify(
        {
            "recognizer_count": len(out),
            "supported_entities": sorted({r.entity for r in recognizers}),
            "recognizers": out,
        }
    )


@app.route("/api/reload", methods=["POST"])
def reload_patterns():
    global recognizers
    try:
        recognizers = build_recognizers()
        return jsonify({"ok": True, "message": "Patterns reloaded."})
    except Exception as e:  # noqa: BLE001
        log.exception("reload failed")
        return jsonify({"ok": False, "message": str(e)}), 500


@app.route("/api/analyze", methods=["POST"])
def analyze():
    if "file" not in request.files:
        return jsonify({"ok": False, "message": "file field is required"}), 400

    f = request.files["file"]
    raw = f.read()
    score_threshold = float(request.form.get("score_threshold", 0.3))
    language = (request.form.get("language") or "auto").strip().lower()
    if language not in ("auto", "ko", "en"):
        language = "auto"

    try:
        text = extract_text(f.filename or "uploaded", raw)
    except Exception as e:  # noqa: BLE001
        log.exception("extract_text failed")
        return jsonify({"ok": False, "message": f"failed to read file: {e}"}), 400

    findings = analyze_text(text, score_threshold, language=language)

    summary: dict[str, int] = {}
    for f_ in findings:
        summary[f_["entity_type"]] = summary.get(f_["entity_type"], 0) + 1

    return jsonify(
        {
            "ok": True,
            "filename": f.filename,
            "char_count": len(text),
            "score_threshold": score_threshold,
            "language": language,
            "summary": summary,
            "findings": findings,
            "text": text,
        }
    )


if __name__ == "__main__":
    host = os.environ.get("HOST", "127.0.0.1")
    port = int(os.environ.get("PORT", "5000"))
    log.info("Starting LITE server (no spaCy/Presidio) on http://%s:%s", host, port)
    app.run(host=host, port=port, debug=False)
